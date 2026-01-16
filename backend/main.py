import os
import shutil
import fitz 
import re
from PIL import Image, ImageOps
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware 
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.core.tesseract_engine import TesseractOCR
from pdf2image import convert_from_path, pdfinfo_from_path # Ensure pdfinfo is imported
import gc


app = FastAPI()

@app.get("/")
def home():
    return {"message": "Sinhala OCR Converter is Working!"}

# Allow frontend access (CORS)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs("uploads", exist_ok=True)

# Initialize Tesseract
ocr_tool = TesseractOCR()

# Extract Images from PDF
def extract_images_from_page(pdf_path, page_index, output_folder):
   
    doc = None
    try:
        doc = fitz.open(pdf_path)
        page = doc[page_index]
        image_list = page.get_images(full=True)
        extracted_paths = []
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            if len(image_bytes) < 3000: continue 
            
            temp_filename = f"{output_folder}/temp_p{page_index}_{img_index}.img"
            with open(temp_filename, "wb") as f:
                f.write(image_bytes)
            
            final_filename = f"{output_folder}/p{page_index}_{img_index}.jpg"
            try:
                with Image.open(temp_filename) as pil_img:
                    pil_img = pil_img.convert("RGB")   
                    pil_img.save(final_filename, "JPEG")
                extracted_paths.append(final_filename)
            except Exception as e:
                print(f"Skipping corrupt image on page {page_index}: {e}")
            finally:
                if os.path.exists(temp_filename):
                    os.remove(temp_filename)

        return extracted_paths
    except Exception as e:
        print(f"Image extraction failed: {e}")
        return []
    finally:
        if doc: doc.close()


# Process Single Page OCR
def process_single_page_ocr(doc, ocr_img, i, file_location, is_pdf):
    
    # Extract and Insert Embedded Images
    if is_pdf:
        extracted = extract_images_from_page(file_location, i, "uploads")
        if extracted:
            if len(extracted) > 1 and i == 0:
                tbl = doc.add_table(rows=1, cols=len(extracted))
                tbl.autofit = True
                for idx, path in enumerate(extracted):
                    try:
                        cell = tbl.cell(0, idx)
                        p = cell.paragraphs[0]
                        r = p.add_run()
                        r.add_picture(path, width=Inches(1.2))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Skipping table image {path}: {e}")


            else:
                for path in extracted:
                    try:
                        doc.add_picture(path, width=Inches(2.5))
                    except Exception as e:
                        print(f"Skipping inline image {path}: {e}")



    # OCR Processing
    try:
        # Extract Text
        lines_data = ocr_tool.extract_text(ocr_img)
        
        # Check for Low Text Scenario
        full_text = " ".join([l['text'] for l in lines_data])
        
        if len(full_text) < 50 and is_pdf:
            print(f"   -> Low text on Page {i+1}. Using Image Fallback.")
            fallback_path = f"uploads/fallback_p{i}.jpg"
            ocr_img.save(fallback_path)
            doc.add_picture(fallback_path, width=Inches(6.0))
            doc.add_page_break()
            return 

        # Add Extracted Text to Document
        for line_item in lines_data:
            line_text = line_item['text']
            alignment = line_item['align']
            
            if not line_text.strip(): continue
            

            # Check for Dotted Line Requirement
            needs_dots = False
            if re.match(r'^(\d+\.|[ivx]+\.)', line_text):
                if line_text.endswith("?") or line_text.endswith("ද?") or len(line_text) > 30:
                    if "......" not in line_text:
                        needs_dots = True

            # Add Paragraph
            p = doc.add_paragraph()
            
            # Set Alignment
            if alignment == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(line_text)
            run.font.name = 'Iskoola Pota'
            

            # Set Font Size and Style
            if alignment == 'center':
                run.bold = True
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(11)

            # Add Dotted Line if Needed
            if needs_dots:
                p_dots = doc.add_paragraph()
                p_dots.add_run(".............................................................................................................")

        doc.add_page_break()

    # Handle OCR Exceptions
    except Exception as e:
        print(f"Error on page {i+1}: {e}")
        p = doc.add_paragraph(f"[OCR Failed: {str(e)}]")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)


@app.post("/convert")
async def convert_file(file: UploadFile = File(...)):
    # Save Uploaded File Locally                     
    file_location = f"uploads/{file.filename}"
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Create Word Document
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    print(f"Processing: {file.filename}")
    is_pdf = file.filename.lower().endswith(".pdf")

    try:
        if is_pdf:
            # PDF Mode
            try:
                
                # Get Number of Pages
                info = pdfinfo_from_path(file_location)
                num_pages = info["Pages"]
            except:
                # Fallback Method
                doc_pdf = fitz.open(file_location)
                num_pages = len(doc_pdf)
                doc_pdf.close()

            print(f"Detected {num_pages} pages. Starting processing...")



            # Process Each Page Individually
            for i in range(num_pages):
                print(f"Processing Page {i+1}/{num_pages}...")
                
                
                # Convert PDF Page to Image
                pages = convert_from_path(
                    file_location, 
                    dpi=450, 
                    first_page=i+1, 
                    last_page=i+1
                )
                
                if pages:
                    current_image = pages[0]
                    
                    # Process OCR for Single Page
                    process_single_page_ocr(doc, current_image, i, file_location, is_pdf=True)
                    
                    # Cleanup
                    del current_image
                    del pages
                    
                    # Garbage Collection
                    gc.collect()

        else:
            # Image Mode
            try:
                # Open Image
                img = Image.open(file_location)
                img = ImageOps.exif_transpose(img) 
                
                # Resize Large Images
                width, height = img.size
                new_size = (width * 3, height * 3)
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                

                # Process OCR for Single Image
                process_single_page_ocr(doc, img, 0, file_location, is_pdf=False)
            except Exception as e:
                return {"error": "Invalid Image File", "details": str(e)}

    # Handle Overall Exceptions
    except Exception as e:
        print(f"Critical Error: {e}")
        return {"error": "Processing Failed", "details": str(e)}

    # Save and Return Document
    output_name = f"uploads/{file.filename}_converted.docx"
    doc.save(output_name)
    return FileResponse(output_name, filename=f"converted_{file.filename}.docx")